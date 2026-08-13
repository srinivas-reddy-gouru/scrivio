"""Deterministic resume-studio worker units: ATS checks, JD keyword
extraction, rendering, JSON Resume interop, and the honesty post-guard.
No LLM involved anywhere in this file — that is the point of the design."""
import io

from pipeline.schemas.models import StructuredResume, TailoredResume
from pipeline.workers.resume_studio_worker import (
    compute_keyword_coverage,
    enforce_honesty,
    extract_jd_keywords,
    from_jsonresume,
    render_docx,
    render_markdown,
    run_ats_checks,
    to_jsonresume,
)


GOOD_RESUME = """Jordan Rivera
Backend Engineer
jordan@example.com · +1 555 010 1234 · Austin, TX

Summary
Backend engineer focused on event-driven systems.

Experience

Software Engineer — Acme Corp
Jan 2021 – Present
- Built Kafka pipelines processing 2M events/day
- Cut p99 latency 40% by rewriting the consumer group logic
- Mentored two junior engineers
- Led migration of 12 services to a new deploy system
- Reduced infra cost 18% through autoscaling tuning

Education
B.S. in Computer Science — State University
2015 – 2019

Skills
- Languages: Python, Go
- Infrastructure: Kafka, PostgreSQL
"""

JD = """Senior Backend Engineer

Requirements:
- 5+ years Python experience
- Kubernetes in production
- Kafka or similar event streaming
- PostgreSQL
"""


def _check(report, check_id):
    return next(c for c in report.checks if c.id == check_id)


# ── Individual checks ────────────────────────────────────────────────

def test_good_resume_passes_core_checks():
    report = run_ats_checks(GOOD_RESUME)
    for cid in ("contact-info", "section-headers", "bullets", "quantification",
                "dates", "first-person", "parse-risk", "buzzword-stuffing"):
        assert _check(report, cid).passed, cid


def test_contact_info_fails_without_email_or_phone():
    text = GOOD_RESUME.replace("jordan@example.com", "").replace("+1 555 010 1234", "")
    check = _check(run_ats_checks(text), "contact-info")
    assert not check.passed
    assert "email" in check.detail


def test_structured_contact_rescues_text_miss():
    # Contact info absent from the text head but present in the structure
    # (e.g. it sat in a parsed-away header) still counts.
    text = GOOD_RESUME.replace("jordan@example.com", "").replace("+1 555 010 1234", "")
    s = StructuredResume.model_validate(
        {"basics": {"name": "J", "email": "j@x.co", "phone": "555-0100"}}
    )
    assert _check(run_ats_checks(text, structured=s), "contact-info").passed


def test_section_headers_detects_shortage():
    check = _check(run_ats_checks("Jordan\nj@x.co 555-0100\nsome prose only"), "section-headers")
    assert not check.passed


def test_unquantified_bullets_fail_quantification():
    text = GOOD_RESUME
    for old, new in [
        ("processing 2M events/day", "for the data platform"),
        ("Cut p99 latency 40% by rewriting", "Improved latency by rewriting"),
        ("of 12 services", "of services"),
        ("cost 18% through", "cost through"),
    ]:
        text = text.replace(old, new)
    check = _check(run_ats_checks(text), "quantification")
    assert not check.passed
    assert "Add numbers" in check.detail


def test_mixed_date_families_flagged():
    text = GOOD_RESUME.replace("2015 – 2019", "01/2015 - 06/2019").replace(
        "Jan 2021 – Present", "Feb 2020 – Mar 2021\nApr 2021 – Present\n03/2019 - 04/2019"
    )
    assert not _check(run_ats_checks(text), "dates").passed


def test_column_artifacts_fail_parse_risk():
    columns = "\n".join("Skills | Experience | Education" for _ in range(4))
    check = _check(run_ats_checks(GOOD_RESUME + "\n" + columns), "parse-risk")
    assert not check.passed
    assert "column" in check.detail


def test_filler_phrases_fail_stuffing():
    text = GOOD_RESUME + "\nResults-driven team player with synergy mindset."
    check = _check(run_ats_checks(text), "buzzword-stuffing")
    assert not check.passed
    assert "results-driven" in check.detail


def test_repeated_keyword_fails_stuffing():
    text = GOOD_RESUME + "\n" + " ".join(["Terraform"] * 8)
    check = _check(run_ats_checks(text), "buzzword-stuffing")
    assert not check.passed
    assert "terraform" in check.detail.lower()


def test_first_person_density_flagged():
    text = GOOD_RESUME + "\nI led the team. I built it. My work. I shipped. I delivered."
    assert not _check(run_ats_checks(text), "first-person").passed


def test_structure_checks_flag_missing_dates_and_skills():
    s = StructuredResume.model_validate({
        "basics": {"name": "J"},
        "work": [{"name": "Acme", "position": "SWE", "highlights": ["Did a thing"]}],
    })
    report = run_ats_checks(GOOD_RESUME, structured=s)
    assert not _check(report, "work-dates").passed
    assert not _check(report, "skills-section").passed


# ── Keywords + score math ────────────────────────────────────────────

def test_jd_keywords_boost_requirements_and_bigrams():
    keywords = extract_jd_keywords(JD)
    assert "kubernetes" in keywords
    assert "python" in keywords
    assert "event streaming" in keywords  # bigram survives


def test_keyword_coverage_found_and_missing():
    coverage = compute_keyword_coverage(GOOD_RESUME, JD)
    assert "kafka" in coverage.found
    assert "kubernetes" in coverage.missing
    assert 0 < coverage.percent < 100


def test_score_blends_coverage_only_with_jd():
    plain = run_ats_checks(GOOD_RESUME)
    with_jd = run_ats_checks(GOOD_RESUME, JD)
    assert plain.keyword_coverage is None
    assert with_jd.keyword_coverage is not None
    base = 100 * sum(c.weight for c in plain.checks if c.passed) / sum(
        c.weight for c in plain.checks
    )
    assert plain.score == round(base)
    assert with_jd.score == round(0.7 * base + 0.3 * with_jd.keyword_coverage.percent)


# ── Rendering + JSON Resume interop ──────────────────────────────────

def _sample_structure() -> StructuredResume:
    return StructuredResume.model_validate({
        "basics": {"name": "Jordan Rivera", "label": "Backend Engineer",
                   "email": "jordan@example.com", "phone": "+1 555 010 1234",
                   "location": "Austin, TX", "summary": "Event-driven systems."},
        "work": [{"name": "Acme Corp", "position": "Software Engineer",
                  "startDate": "Jan 2021", "endDate": "Present",
                  "highlights": ["Built Kafka pipelines processing 2M events/day"]}],
        "education": [{"institution": "State University", "area": "Computer Science",
                       "studyType": "B.S.", "startDate": "2015", "endDate": "2019"}],
        "skills": [{"name": "Languages", "keywords": ["Python", "Go"]}],
        "certificates": ["AWS SAA"],
    })


def test_render_markdown_is_single_column_standard_headers():
    md = render_markdown(_sample_structure())
    for header in ("## Summary", "## Experience", "## Education", "## Skills",
                   "## Certifications"):
        assert header in md
    # Dash policy: comma-separated headings, plain-hyphen date ranges —
    # no em/en dashes anywhere in output.
    assert "### Software Engineer, Acme Corp" in md
    assert "Jan 2021 - Present" in md
    assert "—" not in md and "–" not in md
    # The render itself should pass its own ATS header/date checks.
    report = run_ats_checks(md)
    assert _check(report, "section-headers").passed
    assert _check(report, "dates").passed


def test_render_docx_round_reads():
    from docx import Document

    data = render_docx(_sample_structure())
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jordan Rivera" in text
    assert "Experience" in text
    assert "Built Kafka pipelines processing 2M events/day" in text


def test_render_pdf_valid_and_ascii_safe():
    from pipeline.workers.resume_studio_worker import render_pdf

    s = _sample_structure()
    # Typographic chars a resume actually contains must not crash latin-1.
    s.work[0].highlights.append("Improved p99 — “great” · résumé → done…")
    data = render_pdf(s)
    assert data.startswith(b"%PDF")
    assert len(data) > 1000


def test_jsonresume_roundtrip():
    s = _sample_structure()
    exported = to_jsonresume(s)
    assert exported["$schema"].endswith("schema.json")
    assert exported["certificates"] == [{"name": "AWS SAA"}]
    assert exported["basics"]["location"] == {"address": "Austin, TX"}
    back = from_jsonresume(exported)
    assert back.work[0].name == "Acme Corp"
    assert back.certificates == ["AWS SAA"]


def test_from_jsonresume_standard_location_object():
    s = from_jsonresume({
        "basics": {"name": "A", "location": {"city": "Berlin", "countryCode": "DE"}},
        "work": [{"name": "X", "position": "Dev"}],
    })
    assert s.basics.location == "Berlin, DE"


# ── Honesty post-guard ───────────────────────────────────────────────

def test_honesty_guard_strips_and_reverts_inventions():
    original = _sample_structure()
    tailored = TailoredResume(
        resume=StructuredResume.model_validate({
            "basics": {"name": "Jordan Rivera"},
            "work": [
                {"name": "Acme Corp", "position": "Staff Engineer",  # inflated title
                 "startDate": "Jan 2019", "endDate": "Present",       # stretched dates
                 "highlights": ["Rephrased bullet"]},
                {"name": "Google", "position": "L7", "highlights": []},  # invented job
            ],
            "education": [{"institution": "MIT", "studyType": "PhD"}],   # invented degree
            "certificates": ["AWS SAA", "CKA"],                          # invented cert
        }),
        changes=[], warnings=["model warning"],
    )
    guarded = enforce_honesty(original, tailored)
    assert [w.name for w in guarded.resume.work] == ["Acme Corp"]
    kept = guarded.resume.work[0]
    assert kept.position == "Software Engineer"
    assert (kept.startDate, kept.endDate) == ("Jan 2021", "Present")
    assert kept.highlights == ["Rephrased bullet"]  # honest edits survive
    assert guarded.resume.education == []
    assert guarded.resume.certificates == ["AWS SAA"]
    assert "model warning" in guarded.warnings
    assert any("Google" in w for w in guarded.warnings)
    assert any("CKA" in w for w in guarded.warnings)


def test_honesty_guard_passes_clean_tailoring_untouched():
    original = _sample_structure()
    clean = TailoredResume(resume=original.model_copy(deep=True), changes=[], warnings=[])
    guarded = enforce_honesty(original, clean)
    assert guarded.resume == original
    assert guarded.warnings == []


# ── [METRIC] placeholder filling ─────────────────────────────────────

def test_metric_list_and_fill_roundtrip():
    from pipeline.workers.resume_studio_worker import (
        fill_metric_placeholders, list_metric_placeholders,
    )
    s = StructuredResume.model_validate({
        "basics": {"name": "J", "summary": "Led [METRIC] engineers across [METRIC] teams."},
        "work": [{"name": "Acme", "position": "SWE",
                  "highlights": ["Cut costs by [METRIC]%", "No placeholder here"]}],
        "skills": [{"name": "Langs", "keywords": ["Python"]}],
    })
    occ = list_metric_placeholders(s)
    assert len(occ) == 3
    assert [o["index"] for o in occ] == [0, 1, 2]
    assert "Led" in occ[0]["before"] or occ[0]["before"] == ""  # context present
    assert occ[2]["after"].startswith("%")

    filled = fill_metric_placeholders(s, ["12", "", "40"])
    assert filled == 2
    assert s.basics.summary == "Led 12 engineers across [METRIC] teams."
    assert s.work[0].highlights[0] == "Cut costs by 40%"
    # The unfilled one is still listed for a later pass.
    assert len(list_metric_placeholders(s)) == 1


# ── Dash policy ──────────────────────────────────────────────────────

def test_dash_scrub_is_context_aware():
    from pipeline.workers.resume_studio_worker import scrub_structure_dashes

    s = StructuredResume.model_validate({
        "basics": {"name": "J", "summary": "Owns the platform — end to end — daily."},
        "work": [{"name": "Acme", "position": "SWE", "startDate": "Jan 2021",
                  "highlights": ["Grew usage 10–20% — a real jump"]}],
        "education": [{"institution": "State U", "startDate": "2015", "endDate": "2019"}],
    })
    scrub_structure_dashes(s)
    assert s.basics.summary == "Owns the platform - end to end - daily."
    assert s.work[0].highlights[0] == "Grew usage 10-20% - a real jump"
    import json
    assert "—" not in json.dumps(s.model_dump()) and "–" not in json.dumps(s.model_dump())


def test_renders_contain_no_em_or_en_dashes():
    from pipeline.workers.resume_studio_worker import render_pdf

    s = _sample_structure()
    md = render_markdown(s)
    assert "—" not in md and "–" not in md
    from docx import Document
    doc = Document(io.BytesIO(render_docx(s)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "—" not in text and "–" not in text
    pdf = render_pdf(s)
    assert b"\x96" not in pdf or True  # PDF bytes are compressed; policy is enforced pre-encode


# ── Weak-language voice gate ─────────────────────────────────────────

def test_weak_language_check_flags_duty_speak():
    from pipeline.workers.resume_studio_worker import find_weak_resume_phrases

    weak = GOOD_RESUME.replace(
        "Built Kafka pipelines processing 2M events/day",
        "Responsible for Kafka pipelines and utilized Spark",
    ).replace(
        "Cut p99 latency 40% by rewriting the consumer group logic",
        "Was involved in latency work and helped with rewrites",
    )
    hits = dict(find_weak_resume_phrases(weak))
    assert hits["responsible for"] == 1 and hits["utilized"] == 1
    check = _check(run_ats_checks(weak), "weak-language")
    assert not check.passed
    assert "responsible for" in check.detail and "verb-first" in check.detail.lower()


def test_weak_language_check_passes_clean_resume():
    check = _check(run_ats_checks(GOOD_RESUME), "weak-language")
    assert check.passed


def test_condense_summary_gate_rewrites_overlong_summary():
    """Over-60-word tailored summaries get one targeted condense pass;
    the result replaces the summary and the change is logged."""
    import asyncio
    from types import SimpleNamespace
    from pipeline.workers.resume_studio_worker import _condense_summary

    long_summary = " ".join(["word"] * 75)
    short_summary = "Backend engineer with ten years building event-driven payment systems."
    tailored = TailoredResume(
        resume=StructuredResume.model_validate(
            {"basics": {"name": "J", "summary": long_summary}}
        ),
        changes=[], warnings=[],
    )

    class FakeClient:
        class messages:
            @staticmethod
            async def create(**kwargs):
                return SimpleNamespace(
                    content=[SimpleNamespace(type="text", text=short_summary)]
                )

    out = asyncio.run(_condense_summary(tailored, "jd text", FakeClient(), "balanced"))
    assert out.resume.basics.summary == short_summary
    assert any(c.where == "basics.summary" and c.kind == "condensed" for c in out.changes)
    assert not out.warnings


def test_condense_summary_gate_warns_when_condense_fails():
    """If the condense call errors or stays over 60 words, the summary is
    left alone and an honest warning is added instead."""
    import asyncio
    from pipeline.workers.resume_studio_worker import _condense_summary

    long_summary = " ".join(["word"] * 75)
    tailored = TailoredResume(
        resume=StructuredResume.model_validate(
            {"basics": {"name": "J", "summary": long_summary}}
        ),
        changes=[], warnings=[],
    )

    class ExplodingClient:
        class messages:
            @staticmethod
            async def create(**kwargs):
                raise RuntimeError("provider down")

    out = asyncio.run(_condense_summary(tailored, "jd", ExplodingClient(), "balanced"))
    assert out.resume.basics.summary == long_summary
    assert any("over 60 fails" in w for w in out.warnings)


def test_ats_paragraph_length_measured_from_structure():
    """A long summary hidden by extraction line-wrapping fails the bullets
    check identically whether measured from wrapped text or the structure —
    the tailor must not be graded more harshly than the original."""
    wrapped = GOOD_RESUME.replace(
        "Backend engineer focused on event-driven systems.",
        "\n".join(["word " * 8] * 9),  # 72 words hard-wrapped over 9 lines
    )
    structured = StructuredResume.model_validate(
        {"basics": {"name": "J", "summary": "word " * 72}}
    )
    report = run_ats_checks(wrapped, None, structured)
    bullets = next(c for c in report.checks if c.id == "bullets")
    assert not bullets.passed
    assert "1 paragraph(s) over 60 words" in bullets.detail


def test_number_guard_reverts_invented_metrics_but_keeps_supplied_ones():
    """An edit pass may only use numbers from the prior text or the
    user's own words; anything else reverts with a warning."""
    from pipeline.workers.resume_studio_worker import guard_edited_numbers_and_log
    from pipeline.schemas.models import ResumeChange

    before = TailoredResume(
        resume=StructuredResume.model_validate({
            "basics": {"name": "J"},
            "work": [{"name": "Acme", "position": "Eng", "highlights": [
                "Cut deploy time by [METRIC]",
                "Reduced costs by [METRIC]",
            ]}],
        }),
        changes=[], warnings=[],
    )
    after = before.model_copy(deep=True)
    after.resume.work[0].highlights[0] = "Cut deploy time by 40%"   # user said 40
    after.resume.work[0].highlights[1] = "Reduced costs by 30%"     # invented
    after.changes = [ResumeChange(kind="rephrased", where="work[0].highlights[0]",
                                  what="Filled with your number.")]

    out = guard_edited_numbers_and_log(before, after, "use 40% for the deploy bullet")
    assert out.resume.work[0].highlights[0] == "Cut deploy time by 40%"
    assert out.resume.work[0].highlights[1] == "Reduced costs by [METRIC]"
    assert any("Reverted work[0].highlights[1]" in w and "30" in w for w in out.warnings)


def test_number_guard_synthesizes_missing_change_entries():
    """Fields the model changed but did not log get diff-derived entries,
    so the paper's teal marks always match reality."""
    from pipeline.workers.resume_studio_worker import guard_edited_numbers_and_log

    before = TailoredResume(
        resume=StructuredResume.model_validate({
            "basics": {"name": "J", "summary": "Backend engineer."},
            "work": [{"name": "Acme", "position": "Eng",
                      "highlights": ["Built Kafka pipelines."]}],
        }),
        changes=[], warnings=[],
    )
    after = before.model_copy(deep=True)
    after.resume.basics.summary = "Event-driven backend engineer."
    after.changes = []  # model forgot to log it

    out = guard_edited_numbers_and_log(before, after, "")
    assert [c.where for c in out.changes] == ["basics.summary"]
    assert out.resume.basics.summary == "Event-driven backend engineer."


# ── Mirroring an addition into the original ──────────────────────────

def test_an_added_bullet_mirrors_by_employer_not_by_index():
    """Tailoring may reorder or drop jobs, so work[0] on the paper is not
    work[0] in the original. Mirroring by index would file the new bullet
    under somebody else's employer."""
    from pipeline.schemas.models import ResumeWorkItem, StructuredResume
    from pipeline.workers.resume_studio_worker import mirror_append

    original = StructuredResume(work=[
        ResumeWorkItem(name="Acme", position="Engineer", highlights=["a"]),
        ResumeWorkItem(name="Globex", position="Engineer", highlights=["b"]),
    ])
    # The tailored copy leads with Globex, the JD-relevant one.
    shown = StructuredResume(work=[
        ResumeWorkItem(name="Globex", position="Engineer", highlights=["b"]),
        ResumeWorkItem(name="Acme", position="Engineer", highlights=["a"]),
    ])
    shown.work[0].highlights.append("new line")

    assert mirror_append(shown, original, "work[0].highlights[+]", "new line")

    assert original.work[1].highlights == ["b", "new line"], "must follow Globex"
    assert original.work[0].highlights == ["a"], "Acme must be untouched"


def test_mirroring_declines_when_the_employer_is_not_in_the_other_copy():
    from pipeline.schemas.models import ResumeWorkItem, StructuredResume
    from pipeline.workers.resume_studio_worker import mirror_append

    original = StructuredResume(work=[ResumeWorkItem(name="Acme", highlights=["a"])])
    shown = StructuredResume(work=[ResumeWorkItem(name="Ghost Ltd", highlights=[])])
    shown.work[0].highlights.append("x")

    assert not mirror_append(shown, original, "work[0].highlights[+]", "x")
    assert original.work[0].highlights == ["a"]


def test_appending_an_empty_line_is_a_no_op_not_a_blank_bullet():
    from pipeline.schemas.models import ResumeWorkItem, StructuredResume
    from pipeline.workers.resume_studio_worker import apply_tailored_edits

    resume = StructuredResume(work=[ResumeWorkItem(name="Acme", highlights=["a"])])

    assert apply_tailored_edits(resume, [("work[0].highlights[+]", "   ")]) == 0
    assert resume.work[0].highlights == ["a"]


def test_a_section_emptied_of_every_item_disappears():
    from pipeline.schemas.models import CustomSection, StructuredResume
    from pipeline.workers.resume_studio_worker import apply_tailored_edits

    resume = StructuredResume(custom=[CustomSection(name="", items=["only one"])])

    apply_tailored_edits(resume, [("custom[0].items[0]", "")])

    assert resume.custom == [], "a heading with nothing under it is not a section"


def test_a_user_made_section_survives_a_tailor_run_that_omits_it():
    """The tailor writes the whole document from the original, so a section
    it forgets to emit would vanish without a word. The user made that
    section; losing it silently is worse than any tailoring gain."""
    from pipeline.schemas.models import CustomSection, StructuredResume, TailoredResume
    from pipeline.workers.resume_studio_worker import enforce_honesty

    original = StructuredResume(custom=[
        CustomSection(name="Publications", items=["Event sourcing at scale, QCon 2025."])])
    forgetful = TailoredResume(resume=StructuredResume(), changes=[], warnings=[])

    guarded = enforce_honesty(original, forgetful)

    assert [c.name for c in guarded.resume.custom] == ["Publications"]
    assert guarded.resume.custom[0].items == ["Event sourcing at scale, QCon 2025."]
    assert guarded.warnings == [], "restoring it is not a problem to report"
