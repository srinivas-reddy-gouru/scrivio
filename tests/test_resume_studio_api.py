"""Resume studio API: create (paste/file/JSON Resume/saved job target),
tailor with the honesty flow, downloads in three formats, CRUD."""
import base64
import io
import json
import tempfile
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from api import server
from main import MockAnthropicClient


RESUME_TEXT = """Jordan Rivera
Backend Engineer
jordan@example.com · +1 555 010 1234 · Austin, TX

Summary
Backend engineer focused on event-driven systems.

Experience

Software Engineer — Acme Corp
Jan 2021 – Present
- Built Kafka pipelines processing 2M events/day
- Cut p99 latency 40% by rewriting the consumer group logic
- Mentored two junior engineers

Education
B.S. in Computer Science — State University
2015 – 2019

Skills
- Languages: Python, Go
- Infrastructure: Kafka, PostgreSQL
"""

JD = (
    "Senior Backend Engineer at Acme.\n\nRequirements:\n"
    "- 5+ years Python\n- Kubernetes in production\n- Kafka event streaming\n"
    "- PostgreSQL\n"
)


@pytest.fixture(autouse=True)
def _isolate_output_dir(monkeypatch):
    with tempfile.TemporaryDirectory() as td:
        monkeypatch.setattr(server, "OUTPUT_ROOT", Path(td))
        yield


@pytest.fixture(autouse=True)
def _mock_anthropic_for_api(monkeypatch):
    monkeypatch.setattr(
        server, "_anthropic_client", lambda request: MockAnthropicClient(request)
    )


def _create(client: TestClient, **overrides) -> dict:
    """Create and return the FINISHED doc. Analysis is two-phase: the POST
    returns status='analyzing' with the deterministic report only; under
    TestClient the background LLM phase completes before post() returns,
    so one GET yields the ready doc."""
    body = {"resume_text": RESUME_TEXT}
    body.update(overrides)
    response = client.post("/resumes", json=body)
    assert response.status_code == 200, response.text
    immediate = response.json()
    assert immediate["status"] == "analyzing"
    assert immediate["report"]["score"] > 0  # deterministic report ships NOW
    assert immediate["review"] is None       # LLM phase not in this response
    done = client.get(f"/resumes/{immediate['resume_id']}")
    assert done.status_code == 200
    doc = done.json()
    assert doc["status"] == "ready", doc.get("error")
    return doc


def _tailor(client: TestClient, resume_id: str) -> dict:
    """Kick tailoring and return the finished doc (same two-phase shape)."""
    response = client.post(f"/resumes/{resume_id}/tailor")
    assert response.status_code == 200, response.text
    assert response.json()["tailor_status"] == "tailoring"
    done = client.get(f"/resumes/{resume_id}").json()
    assert done["tailor_status"] == "idle", done.get("tailor_error")
    return done


# ── Create ───────────────────────────────────────────────────────────

def test_create_without_jd_gives_report_and_review():
    client = TestClient(server.app)
    doc = _create(client)
    assert doc["report"]["score"] > 0
    assert doc["report"]["keyword_coverage"] is None
    check_ids = {c["id"] for c in doc["report"]["checks"]}
    assert {"contact-info", "quantification", "work-dates", "skills-section"} <= check_ids
    assert doc["review"]["missing_keywords"] == []  # no JD → no keyword claims
    assert doc["structured"]["basics"]["name"]  # extraction ran
    assert doc["jd_label"] == ""


def test_create_with_pasted_jd_adds_coverage():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    coverage = doc["report"]["keyword_coverage"]
    assert coverage is not None and coverage["percent"] > 0
    assert "kubernetes" in coverage["missing"]
    assert doc["review"]["missing_keywords"] == ["Kubernetes"]
    assert doc["jd_label"] == "pasted JD"


def test_create_with_saved_job_target():
    client = TestClient(server.app)
    profile = client.post("/job-profiles", json={
        "role_title": "Senior Backend Engineer", "company": "Stripe",
        "job_description": JD, "resume_text": RESUME_TEXT,
    })
    assert profile.status_code == 200, profile.text
    profile_id = profile.json()["profile"]["profile_id"]
    doc = _create(client, job_profile_id=profile_id)
    assert doc["jd_label"] == "Senior Backend Engineer @ Stripe"
    assert doc["report"]["keyword_coverage"] is not None


def test_create_via_txt_upload():
    client = TestClient(server.app)
    b64 = base64.b64encode(RESUME_TEXT.encode()).decode()
    doc = _create(
        client, resume_text="", resume_file_b64=b64, resume_filename="resume.txt"
    )
    assert doc["original_text"].startswith("Jordan Rivera")


def test_create_via_jsonresume_upload_skips_extraction(monkeypatch):
    async def boom(*args, **kwargs):  # extraction must NOT be called
        raise AssertionError("extract_resume called for a JSON Resume upload")

    monkeypatch.setattr(server, "extract_resume", boom)
    client = TestClient(server.app)
    jsonresume = {
        "basics": {"name": "Jordan Rivera", "email": "jordan@example.com",
                   "phone": "555-0100"},
        "work": [{"name": "Acme Corp", "position": "Software Engineer",
                  "startDate": "Jan 2021", "endDate": "Present",
                  "highlights": ["Built Kafka pipelines processing 2M events/day"]}],
        "skills": [{"name": "Languages", "keywords": ["Python"]}],
    }
    b64 = base64.b64encode(json.dumps(jsonresume).encode()).decode()
    doc = _create(
        client, resume_text="", resume_file_b64=b64, resume_filename="resume.json"
    )
    assert doc["structured"]["work"][0]["name"] == "Acme Corp"
    assert "# Jordan Rivera" in doc["original_text"]  # rendered from structure


def test_create_validation_errors():
    client = TestClient(server.app)
    assert client.post("/resumes", json={}).status_code == 422
    bad_b64 = client.post("/resumes", json={
        "resume_file_b64": "!!!not-base64!!!", "resume_filename": "r.txt",
    })
    assert bad_b64.status_code == 422
    bad_json = client.post("/resumes", json={
        "resume_file_b64": base64.b64encode(b"{ not json").decode(),
        "resume_filename": "r.json",
    })
    assert bad_json.status_code == 422


# ── Tailor ───────────────────────────────────────────────────────────

def test_tailor_requires_jd():
    client = TestClient(server.app)
    doc = _create(client)  # no JD
    response = client.post(f"/resumes/{doc['resume_id']}/tailor")
    assert response.status_code == 422
    assert "job description" in response.json()["detail"]


def test_tailor_produces_before_after_and_change_log():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    out = _tailor(client, doc["resume_id"])
    tailored = out["tailored"]
    assert tailored["changes"], "change log must not be empty"
    kinds = {c["kind"] for c in tailored["changes"]}
    assert "placeholder" in kinds
    assert any("[METRIC]" in w for w in tailored["warnings"])
    assert any("Cannot honestly claim" in w for w in tailored["warnings"])
    # Honesty guard held: same employer, title, dates as the original.
    work = tailored["resume"]["work"][0]
    assert (work["name"], work["position"]) == ("Acme Corp", "Software Engineer")
    assert (work["startDate"], work["endDate"]) == ("Jan 2021", "Present")
    # Before/after reports both present and both scored.
    assert out["report"]["score"] > 0
    assert out["tailored_report"]["score"] > 0
    assert out["tailored_report"]["keyword_coverage"] is not None


# ── Downloads ────────────────────────────────────────────────────────

def test_download_markdown_and_json():
    client = TestClient(server.app)
    doc = _create(client)
    rid = doc["resume_id"]
    md = client.get(f"/resumes/{rid}/download?fmt=md")
    assert md.status_code == 200
    assert md.headers["content-type"].startswith("text/markdown")
    assert "attachment" in md.headers["content-disposition"]
    assert "## Experience" in md.text
    js = client.get(f"/resumes/{rid}/download?fmt=json")
    assert js.status_code == 200
    exported = js.json()
    assert exported["basics"]["name"] == "Jordan Rivera"
    assert "$schema" in exported


def test_download_docx_round_reads():
    from docx import Document

    client = TestClient(server.app)
    doc = _create(client)
    response = client.get(f"/resumes/{doc['resume_id']}/download?fmt=docx")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats"
    )
    parsed = Document(io.BytesIO(response.content))
    assert any("Jordan Rivera" in p.text for p in parsed.paragraphs)


def test_download_pdf():
    client = TestClient(server.app)
    doc = _create(client)
    response = client.get(f"/resumes/{doc['resume_id']}/download?fmt=pdf")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.content.startswith(b"%PDF")
    assert len(response.content) > 1000


def test_download_tailored_version_gating():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    assert client.get(f"/resumes/{rid}/download?version=tailored").status_code == 404
    _tailor(client, rid)
    tailored_md = client.get(f"/resumes/{rid}/download?version=tailored")
    assert tailored_md.status_code == 200
    assert "[METRIC]" in tailored_md.text
    assert client.get(f"/resumes/{rid}/download?fmt=rtf").status_code == 422
    assert client.get(f"/resumes/{rid}/download?version=draft").status_code == 422


# ── CRUD ─────────────────────────────────────────────────────────────

def test_list_get_delete_roundtrip():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    listing = client.get("/resumes").json()
    assert [item["resume_id"] for item in listing] == [rid]
    assert listing[0]["name"] == "Jordan Rivera"
    assert listing[0]["jd_label"] == "pasted JD"
    assert listing[0]["score"] == doc["report"]["score"]
    fetched = client.get(f"/resumes/{rid}")
    assert fetched.status_code == 200
    assert fetched.json()["resume_id"] == rid
    assert client.delete(f"/resumes/{rid}").json() == {"deleted": rid}
    assert client.get(f"/resumes/{rid}").status_code == 404
    assert client.get("/resumes").json() == []


def test_path_traversal_ids_rejected():
    client = TestClient(server.app)
    for bad in ("invalid~id", "x" * 100, "20250101..000000-abcdef"):
        assert client.get(f"/resumes/{bad}").status_code == 404
        assert client.delete(f"/resumes/{bad}").status_code == 404


# ── Two-phase status flow ────────────────────────────────────────────

def test_review_failure_keeps_checks_and_reports_error(monkeypatch):
    async def boom(*args, **kwargs):
        raise RuntimeError("provider down")

    monkeypatch.setattr(server, "review_resume", boom)
    client = TestClient(server.app)
    response = client.post("/resumes", json={"resume_text": RESUME_TEXT})
    assert response.status_code == 200
    doc = client.get(f"/resumes/{response.json()['resume_id']}").json()
    assert doc["status"] == "error"
    assert "provider" in doc["error"].lower() or "review" in doc["error"].lower()
    assert doc["report"]["score"] > 0     # deterministic report survives
    assert doc["review"] is None


def test_tailor_conflicts_are_guarded():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    # Simulate a tailor already in flight (in prod the background task
    # is still running; under TestClient it finishes instantly, so set
    # the persisted state directly).
    from pipeline.schemas.models import ResumeDoc
    stored = ResumeDoc.model_validate_json(
        server._resume_path(rid).read_text(encoding="utf-8"))
    stored.tailor_status = "tailoring"
    server._save_resume_doc(stored)
    assert client.post(f"/resumes/{rid}/tailor").status_code == 409


# ── [METRIC] fill endpoint ───────────────────────────────────────────

def test_fill_metrics_finishes_the_tailored_resume():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    assert client.post(f"/resumes/{rid}/fill-metrics",
                       json={"values": ["35"]}).status_code == 422  # not tailored yet
    _tailor(client, rid)
    # Mock tailored resume has exactly one [METRIC] (onboarding % highlight).
    r = client.post(f"/resumes/{rid}/fill-metrics", json={"values": ["35"]})
    assert r.status_code == 200, r.text
    out = r.json()
    highlights = out["tailored"]["resume"]["work"][0]["highlights"]
    assert any("by 35%" in h for h in highlights)
    assert not any("[METRIC]" in h for h in highlights)
    # Downloads now come out finished.
    md = client.get(f"/resumes/{rid}/download?fmt=md&version=tailored")
    assert "[METRIC]" not in md.text
    assert "by 35%" in md.text


# ── Edit, coach, instructed edit, undo ───────────────────────────────

def test_manual_edit_updates_text_reruns_checks_and_banks_undo():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    assert client.post(f"/resumes/{rid}/edit-tailored",
                       json={"edits": [{"path": "basics.summary", "value": "x"}]},
                       ).status_code == 422  # not tailored yet
    doc = _tailor(client, rid)
    before_summary = doc["tailored"]["resume"]["basics"]["summary"]
    r = client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "basics.summary", "value": "Hands-on backend engineer for event-driven payment systems."},
    ]})
    assert r.status_code == 200, r.text
    out = r.json()
    assert out["tailored"]["resume"]["basics"]["summary"].startswith("Hands-on backend")
    assert out["tailored_report"] is not None
    assert len(out["tailored_history"]) == 1
    assert out["tailored_history"][0]["resume"]["basics"]["summary"] == before_summary
    # Unknown paths are rejected, not silently dropped.
    bad = client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "work[0].name", "value": "FakeCorp"},
    ]})
    assert bad.status_code == 422
    assert "not editable" in bad.json()["detail"]


def test_manual_edit_empty_value_deletes_a_bullet():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    doc = _tailor(client, rid)
    n_before = len(doc["tailored"]["resume"]["work"][0]["highlights"])
    r = client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "work[0].highlights[0]", "value": ""},
    ]})
    assert r.status_code == 200, r.text
    assert len(r.json()["tailored"]["resume"]["work"][0]["highlights"]) == n_before - 1


def test_advise_answers_without_mutating_the_doc():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    _tailor(client, rid)
    r = client.post(f"/resumes/{rid}/advise", json={
        "question": "What number should go in the deploy-time bullet?",
        "history": [{"role": "user", "content": "hi"},
                    {"role": "assistant", "content": "hello"}],
    })
    assert r.status_code == 200, r.text
    assert "pipeline duration" in r.json()["answer"]
    # Read-only: no undo entry, doc untouched.
    assert client.get(f"/resumes/{rid}").json()["tailored_history"] == []
    assert client.post(f"/resumes/{rid}/advise",
                       json={"question": "  "}).status_code == 422


def test_request_edit_applies_instruction_behind_honesty_guard():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    doc = _tailor(client, rid)
    r = client.post(f"/resumes/{rid}/request-edit",
                    json={"instruction": "Lead the summary with Kafka experience."})
    assert r.status_code == 200, r.text
    out = r.json()
    # Honesty spine holds through instructed edits: employers unchanged.
    orig_names = [w["name"] for w in doc["structured"]["work"]]
    assert [w["name"] for w in out["tailored"]["resume"]["work"]] == orig_names
    assert len(out["tailored_history"]) == 1
    assert client.post(f"/resumes/{rid}/request-edit",
                       json={"instruction": ""}).status_code == 422


def test_undo_walks_back_through_the_history_stack():
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    doc = _tailor(client, rid)
    original_summary = doc["tailored"]["resume"]["basics"]["summary"]
    client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "basics.summary", "value": "First edit."}]})
    client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "basics.summary", "value": "Second edit."}]})
    r = client.post(f"/resumes/{rid}/undo-tailored")
    assert r.status_code == 200, r.text
    assert r.json()["tailored"]["resume"]["basics"]["summary"] == "First edit."
    r = client.post(f"/resumes/{rid}/undo-tailored")
    assert r.json()["tailored"]["resume"]["basics"]["summary"] == original_summary
    assert client.post(f"/resumes/{rid}/undo-tailored").status_code == 422  # stack empty


def test_request_edit_from_conversation_with_empty_instruction():
    """An empty instruction with coach history applies what was
    recommended; empty instruction with no history is still rejected."""
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    _tailor(client, rid)
    assert client.post(f"/resumes/{rid}/request-edit",
                       json={"instruction": ""}).status_code == 422
    r = client.post(f"/resumes/{rid}/request-edit", json={
        "instruction": "",
        "history": [
            {"role": "user", "content": "What metric fits the deploy bullet?"},
            {"role": "assistant", "content": "Use before/after pipeline minutes."},
        ],
    })
    assert r.status_code == 200, r.text
    assert len(r.json()["tailored_history"]) == 1


def test_status_text_in_warnings_is_moved_to_note(monkeypatch):
    """A "no changes this pass" explanation must never persist as an
    honesty note; the endpoint moves it to the transient note field."""
    from pipeline.schemas.models import TailoredResume as TR

    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    doc = _tailor(client, rid)

    async def fake_edit(**kwargs):
        current = TR.model_validate(doc["tailored"])
        current.warnings = ["NO CHANGES MADE THIS PASS: everything needs your facts."]
        return current

    monkeypatch.setattr(server, "edit_resume_by_instruction", fake_edit)
    r = client.post(f"/resumes/{rid}/request-edit", json={"instruction": "fix all notes"})
    assert r.status_code == 200, r.text
    out = r.json()
    assert out["tailored"]["warnings"] == []
    assert out["tailored"]["note"].startswith("NO CHANGES")


def test_change_log_is_cumulative_across_edit_passes():
    """A pass appends its change entries; it never replaces the log.
    The 'what changed' history survives every subsequent fix pass."""
    client = TestClient(server.app)
    doc = _create(client, jd_text=JD)
    rid = doc["resume_id"]
    doc = _tailor(client, rid)
    n_tailor = len(doc["tailored"]["changes"])
    assert n_tailor > 0

    r = client.post(f"/resumes/{rid}/request-edit",
                    json={"instruction": "Lead the summary with Kafka."})
    assert r.status_code == 200, r.text
    n_after = len(r.json()["tailored"]["changes"])
    assert n_after >= n_tailor  # history kept, pass entries appended

    # Manual edits are logged too, so they get teal marks and history.
    r2 = client.post(f"/resumes/{rid}/edit-tailored", json={"edits": [
        {"path": "basics.summary", "value": "My own words."}]})
    changes = r2.json()["tailored"]["changes"]
    assert len(changes) == n_after + 1
    assert changes[-1]["where"] == "basics.summary"
    assert "Edited by you" in changes[-1]["what"]


# ── Adding to the resume ─────────────────────────────────────────────
# The studio's whole spine is that a model may not put anything new on a
# resume. The person whose resume it is has the opposite right, and these
# tests pin the line between those two facts.


def _add(client: TestClient, rid: str, **body) -> dict:
    response = client.post(f"/resumes/{rid}/add", json=body)
    assert response.status_code == 200, response.text
    return response.json()


def test_added_bullet_lands_on_the_paper_and_in_the_original():
    """Writing only to the tailored copy would lose the bullet at the next
    tailor run, which rewrites from the original structure."""
    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    before = len(doc["tailored"]["resume"]["work"][0]["highlights"])

    out = _add(client, doc["resume_id"], kind="bullet", parent="work[0]",
               text="Ran the incident review after every production outage.")

    assert len(out["tailored"]["resume"]["work"][0]["highlights"]) == before + 1
    assert any("incident review" in h for h in out["structured"]["work"][0]["highlights"]), \
        "the addition must reach the original, not just the tailored copy"
    assert out["tailored"]["changes"][-1]["what"] == "Added by you."


def test_the_honesty_guard_keeps_a_job_the_user_added():
    """The guard exists to stop the MODEL inventing an employer. A job the
    user typed is not an invention, and deleting it would be the studio
    calling its own user a liar."""
    from pipeline.workers.resume_studio_worker import enforce_honesty
    from pipeline.schemas.models import StructuredResume, TailoredResume

    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    out = _add(client, doc["resume_id"], kind="work", fields={
        "name": "Initech", "position": "Software Engineer",
        "startDate": "Jan 2014", "endDate": "Jul 2015",
        "highlights": ["Maintained the nightly reconciliation batch."],
    })

    original = StructuredResume.model_validate(out["structured"])
    tailored = TailoredResume.model_validate(out["tailored"])
    guarded = enforce_honesty(original, tailored)

    assert "Initech" in [w.name for w in guarded.resume.work]
    assert not any("Initech" in w for w in guarded.warnings)


def test_the_honesty_guard_still_strips_an_employer_the_model_invents():
    from pipeline.workers.resume_studio_worker import enforce_honesty
    from pipeline.schemas.models import StructuredResume, TailoredResume

    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    original = StructuredResume.model_validate(doc["structured"])
    tailored = TailoredResume.model_validate(doc["tailored"])
    tailored.resume.work.append(
        tailored.resume.work[0].model_copy(deep=True, update={"name": "Umbrella Corp"}))

    guarded = enforce_honesty(original, tailored)

    assert "Umbrella Corp" not in [w.name for w in guarded.resume.work]
    assert any("Umbrella Corp" in w for w in guarded.warnings)


def test_a_user_made_section_reaches_every_download_format():
    import docx
    from pipeline.schemas.models import StructuredResume
    from pipeline.workers.resume_studio_worker import (
        render_docx, render_markdown, render_pdf, to_jsonresume, from_jsonresume,
    )
    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    out = _add(client, doc["resume_id"], kind="custom", fields={
        "name": "Publications", "items": ["Event sourcing at scale, QCon 2025."]})
    out = _add(client, doc["resume_id"], kind="bullet", parent="custom[0]",
               text="Idempotent consumers, ACM Queue 2024.")
    resume = StructuredResume.model_validate(out["tailored"]["resume"])

    assert resume.custom[0].items == [
        "Event sourcing at scale, QCon 2025.", "Idempotent consumers, ACM Queue 2024."]
    assert "## Publications" in render_markdown(resume)
    assert any(p.text == "Publications"
               for p in docx.Document(io.BytesIO(render_docx(resume))).paragraphs)
    assert len(render_pdf(resume)) > 1000
    # JSON Resume has no field for these, so they ride in a namespaced key
    # that other importers ignore and this one round trips.
    exported = to_jsonresume(resume)
    assert exported["x-custom-sections"][0]["name"] == "Publications"
    assert from_jsonresume(exported).custom[0].items == resume.custom[0].items


def test_the_tailor_may_not_invent_a_whole_section():
    from pipeline.workers.resume_studio_worker import enforce_honesty
    from pipeline.schemas.models import CustomSection, StructuredResume, TailoredResume

    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    original = StructuredResume.model_validate(doc["structured"])
    tailored = TailoredResume.model_validate(doc["tailored"])
    tailored.resume.custom.append(
        CustomSection(name="Patents", items=["US1234567, distributed ledgers."]))

    guarded = enforce_honesty(original, tailored)

    assert guarded.resume.custom == []
    assert any("Patents" in w for w in guarded.warnings)


def test_removing_an_entry_drops_it_from_both_copies():
    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    _add(client, doc["resume_id"], kind="custom", fields={
        "name": "Volunteering", "items": ["Mentor at a local code club."]})

    response = client.post(f"/resumes/{doc['resume_id']}/remove-entry",
                           json={"path": "custom[0]"})

    assert response.status_code == 200, response.text
    out = response.json()
    assert out["tailored"]["resume"]["custom"] == []
    assert out["structured"]["custom"] == []


def test_additions_are_undoable_like_any_other_change():
    client = TestClient(server.app)
    doc = _tailor(client, _create(client, jd_text=JD)["resume_id"])
    before = len(doc["tailored"]["resume"]["work"][0]["highlights"])
    _add(client, doc["resume_id"], kind="bullet", parent="work[0]", text="A line I regret.")

    undone = client.post(f"/resumes/{doc['resume_id']}/undo-tailored").json()

    assert len(undone["tailored"]["resume"]["work"][0]["highlights"]) == before


def test_empty_and_unaddressable_additions_are_refused():
    client = TestClient(server.app)
    rid = _tailor(client, _create(client, jd_text=JD)["resume_id"])["resume_id"]
    for body in (
        {"kind": "bullet", "parent": "work[0]", "text": "   "},
        {"kind": "bullet", "parent": "work[99]", "text": "nowhere to go"},
        {"kind": "bullet", "parent": "basics", "text": "not a list"},
        {"kind": "work", "fields": {}},
    ):
        assert client.post(f"/resumes/{rid}/add", json=body).status_code == 422, body


def test_appends_are_refused_on_the_edit_endpoint():
    """One road in for additions. edit-tailored writes only to the tailored
    copy, so an append accepted there would vanish at the next tailor."""
    client = TestClient(server.app)
    rid = _tailor(client, _create(client, jd_text=JD)["resume_id"])["resume_id"]

    response = client.post(f"/resumes/{rid}/edit-tailored", json={
        "edits": [{"path": "work[0].highlights[+]", "value": "sneaking in"}]})

    assert response.status_code == 422
    assert "/add" in response.json()["detail"]
