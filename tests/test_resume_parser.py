import io

import pytest

from pipeline.workers.resume_parser import ResumeParseError, parse_resume


RESUME_TEXT = (
    "Srinivas Reddy — Senior Backend Engineer\n"
    "Experience: built Kafka pipelines processing 2M events/day at Acme. "
    "Led a team of 4. Java, Spring Boot, PostgreSQL, AWS."
)


def _pdf_bytes(text: str) -> bytes:
    """Minimal one-page PDF with real extractable text via pypdf's writer."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    # pypdf can't draw text; use reportlab-free approach: embed via
    # docx-style path is unavailable, so build with a simple content stream.
    import pypdf.generic as g

    stream = f"BT /F1 12 Tf 72 720 Td ({text[:200]}) Tj ET".encode("latin-1", "replace")
    content = g.StreamObject()
    content.set_data(stream)
    page[g.NameObject("/Contents")] = writer._add_object(content)
    page[g.NameObject("/Resources")] = g.DictionaryObject({
        g.NameObject("/Font"): g.DictionaryObject({
            g.NameObject("/F1"): g.DictionaryObject({
                g.NameObject("/Type"): g.NameObject("/Font"),
                g.NameObject("/Subtype"): g.NameObject("/Type1"),
                g.NameObject("/BaseFont"): g.NameObject("/Helvetica"),
            })
        })
    })
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _docx_bytes(text: str) -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Kafka, Spring Boot"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_parse_txt() -> None:
    assert "Kafka pipelines" in parse_resume(RESUME_TEXT.encode(), "resume.txt")


def test_parse_docx_includes_tables() -> None:
    text = parse_resume(_docx_bytes(RESUME_TEXT), "resume.docx")
    assert "Kafka pipelines" in text
    assert "Skills | Kafka, Spring Boot" in text


def test_parse_pdf() -> None:
    text = parse_resume(_pdf_bytes(RESUME_TEXT), "resume.pdf")
    assert "Senior Backend Engineer" in text


def test_scanned_pdf_actionable_error() -> None:
    """An image-only PDF extracts ~no text → actionable message."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    with pytest.raises(ResumeParseError, match="scanned image"):
        parse_resume(buf.getvalue(), "scan.pdf")


def test_unsupported_extension() -> None:
    with pytest.raises(ResumeParseError, match="Unsupported file type"):
        parse_resume(b"binarystuff" * 20, "resume.pages")


def test_empty_file() -> None:
    with pytest.raises(ResumeParseError, match="empty"):
        parse_resume(b"", "resume.pdf")


def test_size_cap() -> None:
    huge = ("word " * 20_000).encode()
    assert len(parse_resume(huge, "resume.txt")) <= 30_000
