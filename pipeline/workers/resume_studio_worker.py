"""Resume studio: structured extraction, deterministic ATS checks,
recruiter-lens review, and honest JD tailoring.

Design (from the Reactive Resume / JSON Resume research): the resume is
extracted into a JSON Resume-shaped structure ONCE, and everything else —
checks, tailoring, the change log, Markdown/DOCX/JSON exports — operates
on that structure, never on a text blob. The ATS score is a transparent
weighted checklist computed in Python (every number explainable in the
UI); the only LLM judgments are the recruiter review and the tailored
rewrite, and the rewrite is followed by a deterministic honesty guard
(employers, titles, and dates in the output must be a subset of the
original's — a lying model gets its inventions stripped, not shipped).
"""
from __future__ import annotations

import io
import re
from collections import Counter
from collections.abc import Sequence

from pipeline.model_config import get_model
from pipeline.workers.citation_utils import scrub_em_dashes
from pipeline.prompt_loader import load_prompt
from pipeline.schemas.models import (
    AtsCheck,
    AtsReport,
    KeywordCoverage,
    ResumeBasics,
    ResumeChange,
    ResumeEducationItem,
    ResumeProject,
    ResumeReview,
    ResumeSkill,
    ResumeWorkItem,
    StructuredResume,
    TailoredResume,
)

_EXTRACTOR_PROMPT = load_prompt("resume_extractor_v1.txt")
_REVIEWER_PROMPT = load_prompt("resume_reviewer_v1.txt")
_TAILOR_PROMPT = load_prompt("resume_tailor_v1.txt")
_ADVISOR_PROMPT = load_prompt("resume_advisor_v1.txt")
_EDITOR_PROMPT = load_prompt("resume_editor_v1.txt")

_EXTRACTION_TOOL: dict = {
    "name": "submit_resume_extraction",
    "description": "Submit the resume mapped faithfully into the JSON Resume structure.",
    "input_schema": StructuredResume.model_json_schema(),
}
_REVIEW_TOOL: dict = {
    "name": "submit_resume_review",
    "description": "Submit the recruiter's review of this resume.",
    "input_schema": ResumeReview.model_json_schema(),
}
_TAILOR_TOOL: dict = {
    "name": "submit_tailored_resume",
    "description": "Submit the tailored resume structure with its change log and honesty warnings.",
    "input_schema": TailoredResume.model_json_schema(),
}


# ── LLM calls ────────────────────────────────────────────────────────────────

async def extract_resume(text: str, client, preset: str = "balanced") -> StructuredResume:
    """Faithful text → structure mapping. No improvement, no invention."""
    response = await client.messages.create(
        model=get_model("resume_extract", preset),
        max_tokens=8192,
        system=_EXTRACTOR_PROMPT,
        tools=[_EXTRACTION_TOOL],
        tool_choice={"type": "tool", "name": "submit_resume_extraction"},
        messages=[{"role": "user", "content": f"resume_text:\n{text}"}],
    )
    tool_use = next(b for b in response.content if b.type == "tool_use")
    return scrub_structure_dashes(StructuredResume.model_validate(tool_use.input))


async def review_resume(
    text: str, jd_text: str, client, preset: str = "balanced"
) -> ResumeReview:
    user_content = (
        f"resume:\n{text}\n\n"
        f"job_description:\n{jd_text or '(none provided — review the resume on its own merits; leave missing_keywords empty)'}"
    )
    response = await client.messages.create(
        model=get_model("resume_review", preset),
        max_tokens=3072,
        system=_REVIEWER_PROMPT,
        tools=[_REVIEW_TOOL],
        tool_choice={"type": "tool", "name": "submit_resume_review"},
        messages=[{"role": "user", "content": user_content}],
    )
    tool_use = next(b for b in response.content if b.type == "tool_use")
    return ResumeReview.model_validate(tool_use.input)


async def tailor_resume(
    *,
    structured: StructuredResume,
    jd_text: str,
    review: ResumeReview | None,
    report: AtsReport | None,
    client,
    preset: str = "balanced",
) -> TailoredResume:
    """JD-targeted rewrite of the STRUCTURE. The prompt carries the honesty
    guardrails; enforce_honesty() is the deterministic belt behind them."""
    review_block = review.model_dump_json(indent=1) if review else "(none)"
    failed = [c for c in (report.checks if report else []) if not c.passed]
    checks_block = (
        "\n".join(f"- {c.label}: {c.detail}" for c in failed) if failed else "(all passed)"
    )
    coverage_block = "(no keyword data)"
    if report and report.keyword_coverage:
        coverage_block = "missing keywords: " + (
            ", ".join(report.keyword_coverage.missing) or "(none)"
        )
    user_content = (
        f"current_resume_structure:\n{structured.model_dump_json(indent=1)}\n\n"
        f"job_description:\n{jd_text}\n\n"
        f"reviewer_findings:\n{review_block}\n\n"
        f"failed_ats_checks:\n{checks_block}\n\n"
        f"keyword_coverage:\n{coverage_block}"
    )
    response = await client.messages.create(
        model=get_model("resume_tailor", preset),
        max_tokens=8192,
        system=_TAILOR_PROMPT,
        tools=[_TAILOR_TOOL],
        tool_choice={"type": "tool", "name": "submit_tailored_resume"},
        messages=[{"role": "user", "content": user_content}],
    )
    tool_use = next(b for b in response.content if b.type == "tool_use")
    tailored = TailoredResume.model_validate(tool_use.input)
    tailored = enforce_honesty(structured, tailored)
    scrub_structure_dashes(tailored.resume)
    # Voice gate: the tailor's own prose is held to the weak-language bar.
    # Deterministic and non-destructive — rewriting phrasing in code would
    # risk meaning; a warning puts the fix in the user's hands.
    weak = find_weak_resume_phrases(render_markdown(tailored.resume))
    if weak:
        tailored.warnings.append(
            "Weak phrasing remains: "
            + ", ".join(f"'{p}' ×{n}" for p, n in weak[:4])
            + ". Consider verb-first rewrites ('Led X', 'Cut Y by 30%')."
        )
    # Length gate: a summary over 60 words fails the same scan check the
    # tailor is scored on, so an overrun un-earns its own points. The cap
    # lives in the prompt; this targeted retry is the deterministic belt
    # for when the model overruns it anyway.
    if len(tailored.resume.basics.summary.split()) > 60:
        tailored = await _condense_summary(tailored, jd_text, client, preset)
    return tailored


async def _condense_summary(
    tailored: TailoredResume, jd_text: str, client, preset: str
) -> TailoredResume:
    """Second, single-purpose pass: cut the summary under 60 words without
    adding anything new. Falls back to an honest warning if it fails."""
    n_before = len(tailored.resume.basics.summary.split())
    try:
        response = await client.messages.create(
            model=get_model("resume_tailor", preset),
            max_tokens=400,
            system=(
                "Condense the resume summary you are given to 55 words or "
                "fewer. Keep the claims most relevant to the job description; "
                "cut the weakest ones entirely. Never add a skill, title, or "
                "claim that is not already in the summary. No em or en "
                "dashes. Reply with the condensed summary text only."
            ),
            messages=[{
                "role": "user",
                "content": (
                    f"summary:\n{tailored.resume.basics.summary}\n\n"
                    f"job_description (for relevance ranking only):\n{jd_text[:2000]}"
                ),
            }],
        )
        text = "".join(
            b.text for b in response.content if getattr(b, "type", "") == "text"
        ).strip()
    except Exception:
        text = ""
    if text and len(text.split()) <= 60:
        tailored.resume.basics.summary = text
        scrub_structure_dashes(tailored.resume)
        tailored.changes.append(ResumeChange(
            kind="condensed", where="basics.summary",
            what=(f"Condensed the summary from {n_before} words to "
                  f"{len(text.split())} so it passes the paragraph-length "
                  "scan check instead of failing it."),
        ))
    else:
        tailored.warnings.append(
            f"The summary is {n_before} words; anything over 60 fails the "
            "scan check. Trim it to the strongest JD-relevant claims."
        )
    return tailored


# ── User-authored edits, coach, and instructed edits ────────────────────────

_HL_PATH_RE = re.compile(r"^(work|projects)\[(\d+)\]\.highlights\[(\d+)\]$")
_SUMMARY_PATH_RE = re.compile(r"^work\[(\d+)\]\.summary$")
_DESC_PATH_RE = re.compile(r"^projects\[(\d+)\]\.description$")


def apply_tailored_edits(
    resume: StructuredResume, edits: Sequence[tuple[str, str]]
) -> int:
    """Apply the user's own text edits, addressed by the same where-paths
    the change log uses. Only prose fields are editable — employers,
    titles, dates, and institutions stay byte-identical to the original
    by construction. An empty value deletes a bullet and clears a
    scalar. Returns how many edits landed; unknown paths raise."""
    applied = 0
    for path, value in edits:
        v = value.strip()
        if path == "basics.label":
            resume.basics.label = v; applied += 1; continue
        if path == "basics.summary":
            resume.basics.summary = v; applied += 1; continue
        m = _SUMMARY_PATH_RE.match(path)
        if m:
            i = int(m.group(1))
            if i < len(resume.work):
                resume.work[i].summary = v; applied += 1
            continue
        m = _DESC_PATH_RE.match(path)
        if m:
            i = int(m.group(1))
            if i < len(resume.projects):
                resume.projects[i].description = v; applied += 1
            continue
        m = _HL_PATH_RE.match(path)
        if m:
            kind, i, j = m.group(1), int(m.group(2)), int(m.group(3))
            items = resume.work if kind == "work" else resume.projects
            if i < len(items) and j < len(items[i].highlights):
                items[i].highlights[j] = v; applied += 1
            continue
        raise ValueError(f"Path not editable: {path}")
    # Emptied bullets vanish — after all index-addressed writes, so the
    # paths the client sent stay valid throughout the batch.
    for w in resume.work:
        w.highlights = [h for h in w.highlights if h.strip()]
    for p in resume.projects:
        p.highlights = [h for h in p.highlights if h.strip()]
    scrub_structure_dashes(resume)
    return applied


def build_resume_advice_context(doc) -> str:
    """Compact context block for the coach: the JD, the current resume
    text, and exactly which numbers are still placeholders."""
    resume = doc.tailored.resume if doc.tailored else doc.structured
    parts = []
    if doc.jd_text:
        parts.append(f"job_description (excerpt):\n{doc.jd_text[:2000]}")
    if resume is not None:
        parts.append(f"current_resume:\n{render_markdown(resume)[:3500]}")
        holes = list_metric_placeholders(resume)
        if holes:
            parts.append("unfilled_metric_placeholders:\n" + "\n".join(
                f"- …{h['before'][-50:]} [METRIC] {h['after']}…" for h in holes))
    if doc.tailored and doc.tailored.warnings:
        parts.append("honesty_warnings:\n" + "\n".join(
            f"- {w}" for w in doc.tailored.warnings[:6]))
    return "\n\n".join(parts) or "(no resume on file)"


async def advise_resume(
    *, question: str, history: Sequence[dict], context: str,
    client, preset: str = "balanced",
) -> str:
    """One coaching answer. Read-only: this call never edits the resume,
    so the user can ask freely before committing to anything."""
    messages = [
        {"role": m["role"], "content": m["content"]}
        for m in history if m.get("role") in ("user", "assistant")
    ][-8:]
    messages.append({
        "role": "user",
        "content": f"context:\n{context}\n\nquestion:\n{question}",
    })
    response = await client.messages.create(
        model=get_model("resume_review", preset),
        max_tokens=700,
        system=_ADVISOR_PROMPT,
        messages=messages,
    )
    text = "".join(
        b.text for b in response.content if getattr(b, "type", "") == "text"
    ).strip()
    return scrub_em_dashes(text) if text else (
        "I could not produce an answer just now. Try rephrasing the question."
    )


async def edit_resume_by_instruction(
    *, original: StructuredResume, tailored: TailoredResume,
    jd_text: str, instruction: str, client, preset: str = "balanced",
) -> TailoredResume:
    """Apply ONE user instruction to the tailored resume via the LLM,
    behind the same honesty guard as tailoring: the model proposes, the
    deterministic post-guard disposes."""
    user_content = (
        f"current_tailored_resume:\n{tailored.resume.model_dump_json(indent=1)}\n\n"
        f"job_description (context only):\n{jd_text[:2000]}\n\n"
        f"user_instruction:\n{instruction}"
    )
    response = await client.messages.create(
        model=get_model("resume_tailor", preset),
        max_tokens=8192,
        system=_EDITOR_PROMPT,
        tools=[_TAILOR_TOOL],
        tool_choice={"type": "tool", "name": "submit_tailored_resume"},
        messages=[{"role": "user", "content": user_content}],
    )
    tool_use = next(b for b in response.content if b.type == "tool_use")
    edited = TailoredResume.model_validate(tool_use.input)
    edited = enforce_honesty(original, edited)
    scrub_structure_dashes(edited.resume)
    return edited


# ── Honesty post-guard ──────────────────────────────────────────────────────

def enforce_honesty(
    original: StructuredResume, tailored: TailoredResume
) -> TailoredResume:
    """Deterministic backstop: facts the model may not touch — employers,
    titles, employment dates, schools, degrees, certificates — must be a
    subset of the original's. Inventions are stripped (or reverted) and
    surfaced as warnings, never silently shipped."""
    warnings = list(tailored.warnings)

    orig_by_employer = {w.name: w for w in original.work if w.name}
    kept_work: list[ResumeWorkItem] = []
    for item in tailored.resume.work:
        source = orig_by_employer.get(item.name)
        if source is None:
            warnings.append(
                f"Removed work entry '{item.name or item.position}' — that employer "
                "is not on the original resume."
            )
            continue
        if item.position != source.position:
            warnings.append(
                f"Reverted job title at {item.name} to '{source.position}' — titles "
                "cannot be changed."
            )
            item.position = source.position
        if (item.startDate, item.endDate) != (source.startDate, source.endDate):
            warnings.append(
                f"Reverted employment dates at {item.name} — dates cannot be changed."
            )
            item.startDate, item.endDate = source.startDate, source.endDate
        kept_work.append(item)
    tailored.resume.work = kept_work

    orig_schools = {e.institution for e in original.education if e.institution}
    kept_edu = []
    for edu in tailored.resume.education:
        if edu.institution and edu.institution not in orig_schools:
            warnings.append(
                f"Removed education entry '{edu.institution}' — that institution "
                "is not on the original resume."
            )
            continue
        kept_edu.append(edu)
    tailored.resume.education = kept_edu

    orig_certs = set(original.certificates)
    invented_certs = [c for c in tailored.resume.certificates if c not in orig_certs]
    if invented_certs:
        tailored.resume.certificates = [
            c for c in tailored.resume.certificates if c in orig_certs
        ]
        warnings.append(
            "Removed certificates not on the original resume: "
            + ", ".join(invented_certs)
        )

    tailored.warnings = warnings
    return tailored


# ── Weak-language detection ─────────────────────────────────────────────────
# The resume-genre equivalent of the article pipeline's banned-phrases
# gate: deterministic detection of duty-speak and filler that recruiters
# (and AI-content detectors) flag. Used twice — as an explainable ATS
# check on any resume, and as a post-tailor warning gate so the tailor's
# own output is held to the same bar.

_WEAK_PHRASES = (
    "responsible for", "duties included", "tasked with", "in charge of",
    "worked on", "helped with", "assisted with", "participated in",
    "was involved in", "leveraged", "utilized", "utilizing",
    "proven track record", "passionate about", "seasoned professional",
    "highly motivated", "cutting-edge", "state-of-the-art",
    "best-in-class", "world-class",
)


def find_weak_resume_phrases(text: str) -> list[tuple[str, int]]:
    """Weak/duty-speak phrases with counts, most frequent first."""
    lower = text.lower()
    hits = [(p, lower.count(p)) for p in _WEAK_PHRASES if p in lower]
    return sorted(hits, key=lambda h: -h[1])


# ── Dash policy ─────────────────────────────────────────────────────────────
# No em/en dashes anywhere in resume output (user rule; they also read as
# an AI tell). Context-aware replacement: digit ranges get a bare hyphen
# ("2015-2019"), clause-break em dashes get a spaced hyphen, everything
# is normalized so no double spaces remain. Applied at every chokepoint
# where text enters a structure: extraction, tailoring, JSON import.

_DIGIT_RANGE_DASH_RE = re.compile(r"(?<=\d)\s*[–—]\s*(?=\d)")
_ANY_DASH_RE = re.compile(r"\s*[–—―]\s*")


def _clean_dashes(text: str) -> str:
    if "–" not in text and "—" not in text and "―" not in text:
        return text
    text = _DIGIT_RANGE_DASH_RE.sub("-", text)
    text = _ANY_DASH_RE.sub(" - ", text)
    return re.sub(r" {2,}", " ", text).strip()


def scrub_structure_dashes(s: StructuredResume) -> StructuredResume:
    """Strip em/en dashes from every text field, in place."""
    b = s.basics
    b.name, b.label, b.summary, b.location = (
        _clean_dashes(b.name), _clean_dashes(b.label),
        _clean_dashes(b.summary), _clean_dashes(b.location),
    )
    for w in s.work:
        w.name, w.position = _clean_dashes(w.name), _clean_dashes(w.position)
        w.startDate, w.endDate = _clean_dashes(w.startDate), _clean_dashes(w.endDate)
        w.summary = _clean_dashes(w.summary)
        w.highlights = [_clean_dashes(h) for h in w.highlights]
    for e in s.education:
        e.institution, e.area = _clean_dashes(e.institution), _clean_dashes(e.area)
        e.studyType, e.score = _clean_dashes(e.studyType), _clean_dashes(e.score)
        e.startDate, e.endDate = _clean_dashes(e.startDate), _clean_dashes(e.endDate)
    for sk in s.skills:
        sk.name = _clean_dashes(sk.name)
        sk.keywords = [_clean_dashes(k) for k in sk.keywords]
    for p in s.projects:
        p.name, p.description = _clean_dashes(p.name), _clean_dashes(p.description)
        p.highlights = [_clean_dashes(h) for h in p.highlights]
    s.certificates = [_clean_dashes(c) for c in s.certificates]
    return s


# ── [METRIC] placeholder filling ────────────────────────────────────────────
# The tailor writes [METRIC] instead of inventing numbers; this pair lets
# the user fill their real figures in-app so downloads come out finished.
# CANONICAL TRAVERSAL ORDER — the UI builds its input list with the same
# walk, so occurrence N here is occurrence N on screen: basics.summary,
# then per work item (summary, highlights…), per project (description,
# highlights…), per skill (keywords…), certificates.

METRIC_TOKEN = "[METRIC]"


def _metric_fields(s: StructuredResume):
    """Yield (get, set) accessors for every free-text field that can carry
    placeholders, in canonical order."""
    yield (lambda: s.basics.summary, lambda v: setattr(s.basics, "summary", v))
    for w in s.work:
        yield (lambda w=w: w.summary, lambda v, w=w: setattr(w, "summary", v))
        for i in range(len(w.highlights)):
            yield (lambda w=w, i=i: w.highlights[i],
                   lambda v, w=w, i=i: w.highlights.__setitem__(i, v))
    for p in s.projects:
        yield (lambda p=p: p.description, lambda v, p=p: setattr(p, "description", v))
        for i in range(len(p.highlights)):
            yield (lambda p=p, i=i: p.highlights[i],
                   lambda v, p=p, i=i: p.highlights.__setitem__(i, v))
    for sk in s.skills:
        for i in range(len(sk.keywords)):
            yield (lambda sk=sk, i=i: sk.keywords[i],
                   lambda v, sk=sk, i=i: sk.keywords.__setitem__(i, v))
    for i in range(len(s.certificates)):
        yield (lambda s=s, i=i: s.certificates[i],
               lambda v, s=s, i=i: s.certificates.__setitem__(i, v))


def list_metric_placeholders(s: StructuredResume) -> list[dict]:
    """Every [METRIC] occurrence in canonical order with display context."""
    found: list[dict] = []
    for get, _ in _metric_fields(s):
        text = get()
        start = 0
        while True:
            pos = text.find(METRIC_TOKEN, start)
            if pos == -1:
                break
            found.append({
                "index": len(found),
                "before": text[max(0, pos - 60):pos].lstrip(),
                "after": text[pos + len(METRIC_TOKEN):pos + len(METRIC_TOKEN) + 40].rstrip(),
            })
            start = pos + len(METRIC_TOKEN)
    return found


def fill_metric_placeholders(s: StructuredResume, values: list[str]) -> int:
    """Replace [METRIC] occurrences in canonical order with *values*.
    Empty/whitespace values leave that placeholder in place (the user can
    fill the rest later). Mutates *s*; returns how many were filled."""
    filled = 0
    occurrence = 0
    for get, set_ in _metric_fields(s):
        text = get()
        if METRIC_TOKEN not in text:
            continue
        out: list[str] = []
        start = 0
        while True:
            pos = text.find(METRIC_TOKEN, start)
            if pos == -1:
                out.append(text[start:])
                break
            out.append(text[start:pos])
            value = values[occurrence].strip() if occurrence < len(values) else ""
            if value:
                out.append(value)
                filled += 1
            else:
                out.append(METRIC_TOKEN)
            occurrence += 1
            start = pos + len(METRIC_TOKEN)
        set_("".join(out))
    return filled


# ── JSON Resume interop ─────────────────────────────────────────────────────

def to_jsonresume(s: StructuredResume) -> dict:
    """Export in the JSON Resume open-standard shape (jsonresume.org) —
    importable by Reactive Resume and the wider ecosystem."""
    data = s.model_dump()
    loc = data["basics"].pop("location", "")
    data["basics"]["location"] = {"address": loc} if loc else {}
    data["certificates"] = [{"name": c} for c in data["certificates"]]
    data["$schema"] = (
        "https://raw.githubusercontent.com/jsonresume/resume-schema/master/schema.json"
    )
    return data


def from_jsonresume(data: dict) -> StructuredResume:
    """Import a JSON Resume document, tolerating the standard's optional
    fields and object-shaped location. Unknown sections are ignored."""
    if not isinstance(data, dict):
        raise ValueError("JSON Resume import expects a top-level object.")
    basics_in = data.get("basics") or {}
    loc = basics_in.get("location") or {}
    if isinstance(loc, dict):
        loc_str = ", ".join(
            str(v) for v in (loc.get("city"), loc.get("region"), loc.get("countryCode"))
            if v
        ) or str(loc.get("address") or "")
    else:
        loc_str = str(loc)
    basics = ResumeBasics.model_validate(
        {**{k: v for k, v in basics_in.items() if isinstance(v, str)}, "location": loc_str}
    )
    certs = []
    for c in data.get("certificates") or []:
        if isinstance(c, dict) and c.get("name"):
            certs.append(str(c["name"]))
        elif isinstance(c, str) and c:
            certs.append(c)
    return scrub_structure_dashes(StructuredResume(
        basics=basics,
        work=[ResumeWorkItem.model_validate(w) for w in data.get("work") or []],
        education=[
            ResumeEducationItem.model_validate(e) for e in data.get("education") or []
        ],
        skills=[ResumeSkill.model_validate(k) for k in data.get("skills") or []],
        projects=[ResumeProject.model_validate(p) for p in data.get("projects") or []],
        certificates=certs,
    ))


# ── Rendering (pure functions, no LLM) ──────────────────────────────────────

def _date_range(start: str, end: str) -> str:
    if start and end:
        return f"{start} - {end}"
    return start or end or ""


def render_markdown(s: StructuredResume) -> str:
    """Single-column, standard-headers markdown — the ATS-safe layout."""
    b = s.basics
    lines: list[str] = [f"# {b.name}".rstrip()]
    if b.label:
        lines.append(f"**{b.label}**")
    contact = " · ".join(p for p in (b.email, b.phone, b.location, b.url) if p)
    if contact:
        lines.append(contact)
    if b.summary:
        lines += ["", "## Summary", "", b.summary]
    if s.work:
        lines += ["", "## Experience"]
        for w in s.work:
            heading = ", ".join(p for p in (w.position, w.name) if p)
            lines += ["", f"### {heading}" if heading else "### Role"]
            dates = _date_range(w.startDate, w.endDate)
            if dates:
                lines.append(dates)
            if w.summary:
                lines.append(w.summary)
            lines += [f"- {h}" for h in w.highlights]
    if s.projects:
        lines += ["", "## Projects"]
        for p in s.projects:
            lines += ["", f"### {p.name}" if p.name else "### Project"]
            if p.description:
                lines.append(p.description)
            if p.url:
                lines.append(p.url)
            lines += [f"- {h}" for h in p.highlights]
    if s.education:
        lines += ["", "## Education"]
        for e in s.education:
            degree = " in ".join(p for p in (e.studyType, e.area) if p)
            heading = ", ".join(p for p in (degree, e.institution) if p)
            lines += ["", f"### {heading}" if heading else "### Education"]
            tail = " · ".join(
                p for p in (_date_range(e.startDate, e.endDate), e.score) if p
            )
            if tail:
                lines.append(tail)
    if s.skills:
        lines += ["", "## Skills", ""]
        for sk in s.skills:
            kw = ", ".join(sk.keywords)
            lines.append(f"- **{sk.name}:** {kw}" if sk.name else f"- {kw}")
    if s.certificates:
        lines += ["", "## Certifications", ""]
        lines += [f"- {c}" for c in s.certificates]
    return "\n".join(lines).strip() + "\n"


def render_docx(s: StructuredResume) -> bytes:
    """Clean DOCX from structure: real Heading styles and bullet lists —
    field mapping beats markdown-to-docx heuristics."""
    from docx import Document

    doc = Document()
    b = s.basics
    doc.add_heading(b.name or "Resume", level=0)
    if b.label:
        doc.add_paragraph(b.label)
    contact = " · ".join(p for p in (b.email, b.phone, b.location, b.url) if p)
    if contact:
        doc.add_paragraph(contact)
    if b.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(b.summary)
    if s.work:
        doc.add_heading("Experience", level=1)
        for w in s.work:
            heading = ", ".join(p for p in (w.position, w.name) if p)
            doc.add_heading(heading or "Role", level=2)
            dates = _date_range(w.startDate, w.endDate)
            if dates:
                doc.add_paragraph(dates)
            if w.summary:
                doc.add_paragraph(w.summary)
            for h in w.highlights:
                doc.add_paragraph(h, style="List Bullet")
    if s.projects:
        doc.add_heading("Projects", level=1)
        for p in s.projects:
            doc.add_heading(p.name or "Project", level=2)
            if p.description:
                doc.add_paragraph(p.description)
            for h in p.highlights:
                doc.add_paragraph(h, style="List Bullet")
    if s.education:
        doc.add_heading("Education", level=1)
        for e in s.education:
            degree = " in ".join(p for p in (e.studyType, e.area) if p)
            heading = ", ".join(p for p in (degree, e.institution) if p)
            doc.add_heading(heading or "Education", level=2)
            tail = " · ".join(
                p for p in (_date_range(e.startDate, e.endDate), e.score) if p
            )
            if tail:
                doc.add_paragraph(tail)
    if s.skills:
        doc.add_heading("Skills", level=1)
        for sk in s.skills:
            kw = ", ".join(sk.keywords)
            doc.add_paragraph(f"{sk.name}: {kw}" if sk.name else kw, style="List Bullet")
    if s.certificates:
        doc.add_heading("Certifications", level=1)
        for c in s.certificates:
            doc.add_paragraph(c, style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# PDF core fonts use WinAnsi encoding: 0x95 renders as a real bullet
# glyph and 0x96 as an en dash, so those survive; everything else
# non-encodable maps to a plain equivalent rather than embedding a font.
_PDF_BULLET = "\x95"
_PDF_ENDASH = "\x96"


def _pdf_safe(text: str) -> str:
    """fpdf2's built-in core fonts are WinAnsi; swap typographic characters
    resumes actually contain for renderable equivalents rather than
    embedding a font (keeps the dependency footprint tiny)."""
    replacements = {
        "\u2013": "-", "\u2014": "-",
        "\u2022": _PDF_BULLET, "\u00b7": "|",
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2026": "...", "\u00a0": " ", "\u2192": "->",
    }
    for src_ch, dst in replacements.items():
        text = text.replace(src_ch, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def render_pdf(s: StructuredResume) -> bytes:
    """Single-column, ATS-safe PDF — the format most application portals
    demand. Same field mapping as the DOCX render, core fonts only.
    Typography rules: real bullet glyphs with hanging indents (wrapped
    lines align under the text, not the margin), bold role lines with
    muted dates, ruled section headers."""
    from fpdf import FPDF

    pdf = FPDF(format="letter")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(18, 15, 18)
    pdf.add_page()
    width = pdf.w - pdf.l_margin - pdf.r_margin

    INK, MUTED, ACCENT, RULE = (25, 30, 29), (95, 105, 103), (26, 84, 78), (185, 195, 193)

    def line(text, size=10, style="", color=INK, spacing=1.42, before=0):
        if before:
            pdf.ln(before)
        pdf.set_font("Helvetica", style, size)
        pdf.set_text_color(*color)
        # new_x=LMARGIN: fpdf2's default parks the cursor at the cell's
        # right edge, which truncates every following line to zero width.
        pdf.multi_cell(
            width, size * 0.353 * spacing, _pdf_safe(text),
            new_x="LMARGIN", new_y="NEXT",
        )

    def bullet(text, size=9.5):
        """Real bullet glyph + hanging indent: wrapped lines align under
        the text start, which is what makes a dense role scannable."""
        h = size * 0.353 * 1.42
        pdf.set_font("Helvetica", "", size)
        pdf.set_text_color(*ACCENT)
        pdf.cell(4.5, h, _PDF_BULLET)
        pdf.set_text_color(*INK)
        pdf.multi_cell(width - 4.5, h, _pdf_safe(text),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.6)

    def labeled(label, rest, size=9.5):
        """Bold inline label ("Languages:") followed by wrapped plain text."""
        h = size * 0.353 * 1.42
        pdf.set_font("Helvetica", "B", size)
        label_w = pdf.get_string_width(_pdf_safe(label)) + 1.5
        pdf.set_text_color(*INK)
        pdf.cell(label_w, h, _pdf_safe(label))
        pdf.set_font("Helvetica", "", size)
        pdf.multi_cell(width - label_w, h, _pdf_safe(rest),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.6)

    def section(title):
        pdf.ln(3.2)
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*ACCENT)
        pdf.set_char_spacing(0.6)
        pdf.multi_cell(width, 5, title.upper(), new_x="LMARGIN", new_y="NEXT")
        pdf.set_char_spacing(0)
        y = pdf.get_y() + 0.4
        pdf.set_draw_color(*RULE)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, y, pdf.l_margin + width, y)
        pdf.ln(2.4)

    b = s.basics
    line(b.name or "Resume", size=20, style="B")
    if b.label:
        line(b.label, size=11, color=(60, 72, 70), before=0.6)
    contact = "  |  ".join(p for p in (b.email, b.phone, b.location, b.url) if p)
    if contact:
        line(contact, size=9, color=MUTED, before=0.8)
    if b.summary:
        section("Summary")
        line(b.summary, size=9.7)
    if s.work:
        section("Experience")
        for i, w in enumerate(s.work):
            heading = ", ".join(p for p in (w.position, w.name) if p)
            line(heading or "Role", style="B", size=10.5, before=2.6 if i else 0)
            dates = _date_range(w.startDate, w.endDate)
            if dates:
                line(dates, size=8.8, color=MUTED, before=0.3)
            pdf.ln(1)
            if w.summary:
                line(w.summary, size=9.5)
                pdf.ln(0.6)
            for h in w.highlights:
                bullet(h)
    if s.projects:
        section("Projects")
        for i, proj in enumerate(s.projects):
            line(proj.name or "Project", style="B", size=10.5, before=2.6 if i else 0)
            if proj.description:
                line(proj.description, size=9.5)
            if proj.url:
                line(proj.url, size=8.8, color=MUTED)
            pdf.ln(0.6)
            for h in proj.highlights:
                bullet(h)
    if s.education:
        section("Education")
        for i, e in enumerate(s.education):
            degree = " in ".join(p for p in (e.studyType, e.area) if p)
            heading = ", ".join(p for p in (degree, e.institution) if p)
            line(heading or "Education", style="B", size=10.5, before=2 if i else 0)
            tail = "  |  ".join(
                p for p in (_date_range(e.startDate, e.endDate), e.score) if p
            )
            if tail:
                line(tail, size=8.8, color=MUTED, before=0.3)
    if s.skills:
        section("Skills")
        for sk in s.skills:
            kw = ", ".join(sk.keywords)
            if sk.name:
                labeled(f"{sk.name}:", kw)
            else:
                line(kw, size=9.5)
    if s.certificates:
        section("Certifications")
        for c in s.certificates:
            bullet(c)
    return bytes(pdf.output())


# ── Deterministic ATS checks ────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+?\d[\d ()\-.]{7,}\d)")
_BULLET_RE = re.compile(r"^\s*[-•*–▪●◦]\s+")
_HEADER_WORDS = (
    "experience", "employment", "work history",
    "education",
    "skills",
    "summary", "profile", "objective",
    "projects",
    "certifications", "certificates", "licenses",
)
_MONTH_DATE_RE = re.compile(
    r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}",
    re.IGNORECASE,
)
_NUMERIC_DATE_RE = re.compile(r"\b\d{1,2}/\d{4}\b")
_YEAR_RANGE_RE = re.compile(r"\b(?:19|20)\d{2}\s*[-–—]", re.IGNORECASE)
# "I" must stay case-sensitive (lowercase "i" would catch list markers),
# but sentence-start "My"/"Me" count.
_FIRST_PERSON_RE = re.compile(r"\b(?:I|[Mm]e|[Mm]y)\b")
_QUANT_RE = re.compile(r"\d|%|\$|€|£")
_MULTISPACE_RE = re.compile(r"\S\s{3,}\S")

_FILLER_PHRASES = (
    "synergy", "synergies", "go-getter", "go getter", "results-driven",
    "results driven", "think outside the box", "hard worker", "team player",
    "self-starter", "detail-oriented", "guru", "ninja", "rockstar",
)

_STOPWORDS = frozenset("""
a an and are as at be been but by for from has have he her his i if in into is
it its me my nor not of on or our out she so that the their them then there
these they this to was we were what when where which while who will with you
your years experience work working strong ability able etc using use team
role responsibilities requirements qualifications preferred must plus new
including knowledge skills skill similar related relevant s t re ll d
""".split())

# Tokens too generic to flag as "stuffed" even at high counts.
_STUFFING_EXEMPT = frozenset((
    "experience", "engineering", "software", "development", "team", "data",
    "product", "design", "management", "project", "systems", "services",
))


def _score(checks: list[AtsCheck], coverage: KeywordCoverage | None) -> int:
    total = sum(c.weight for c in checks) or 1
    passed = sum(c.weight for c in checks if c.passed)
    base = 100 * passed / total
    if coverage is not None:
        return round(0.7 * base + 0.3 * coverage.percent)
    return round(base)


def run_ats_checks(
    text: str,
    jd_text: str | None = None,
    structured: StructuredResume | None = None,
) -> AtsReport:
    """Weighted, explainable checklist — pure Python, unit-testable.
    Score = weighted passed checks normalized to 100; with a JD it becomes
    70% checks + 30% keyword coverage."""
    lines = [ln.rstrip() for ln in text.splitlines()]
    nonempty = [ln for ln in lines if ln.strip()]
    words = text.split()
    checks: list[AtsCheck] = []

    def add(id_: str, label: str, passed: bool, weight: int, detail: str) -> None:
        checks.append(
            AtsCheck(id=id_, label=label, passed=passed, weight=weight, detail=detail)
        )

    # contact-info — email + phone near the top (ATS field parsing reads there).
    head = "\n".join(nonempty[:15])
    has_email = bool(_EMAIL_RE.search(head)) or (
        structured is not None and bool(structured.basics.email)
    )
    has_phone = bool(_PHONE_RE.search(head)) or (
        structured is not None and bool(structured.basics.phone)
    )
    add(
        "contact-info", "Contact info up top", has_email and has_phone, 15,
        "Email and phone found near the top."
        if has_email and has_phone
        else "Missing "
        + " and ".join(p for p, ok in (("an email", has_email), ("a phone number", has_phone)) if not ok)
        + " in the top lines — ATS field parsers read contact details there.",
    )

    # section-headers — ≥3 standard headers at line starts.
    found_headers = set()
    for ln in nonempty:
        stripped = ln.strip().strip("#*").strip().lower().rstrip(":")
        for h in _HEADER_WORDS:
            if stripped == h or stripped.startswith(h + " "):
                found_headers.add(h.split()[0])
    add(
        "section-headers", "Standard section headers", len(found_headers) >= 3, 15,
        f"Recognized sections: {', '.join(sorted(found_headers)) or 'none'}. "
        + ("Good — ATS segmentation keys off these."
           if len(found_headers) >= 3
           else "Fewer than 3 standard headers (Experience, Education, Skills, Summary, Projects, Certifications) — ATS segmentation may misfile content."),
    )

    # length — 350-1100 words.
    wc = len(words)
    add(
        "length", "Length in range", 350 <= wc <= 1100, 10,
        f"{wc} words. "
        + ("Within the 350–1100 sweet spot."
           if 350 <= wc <= 1100
           else "Under 350 reads thin — add substance." if wc < 350
           else "Over 1100 words (2+ pages) — condense to the most relevant material."),
    )

    # bullets — enough bullet lines, no wall-of-text paragraphs. When the
    # structure is known, paragraph length comes from the STRUCTURE's prose
    # fields, not the text: extraction line-wrapping would hide a long
    # summary in the original while the tailored rendering shows it as one
    # line, grading the tailor more harshly than the source it improved.
    # An ATS parses the paragraph, not the line breaks.
    bullet_lines = [ln for ln in nonempty if _BULLET_RE.match(ln)]
    if structured is not None:
        prose = [structured.basics.summary]
        prose += [w.summary for w in structured.work]
        prose += [p.description for p in structured.projects]
        long_paras = [p for p in prose if len(p.split()) > 60]
    else:
        long_paras = [
            ln for ln in nonempty
            if not _BULLET_RE.match(ln) and len(ln.split()) > 60
        ]
    bullets_ok = len(bullet_lines) >= 5 and not long_paras
    add(
        "bullets", "Bullet-driven experience", bullets_ok, 10,
        f"{len(bullet_lines)} bullet lines, {len(long_paras)} paragraph(s) over 60 words. "
        + ("Scannable." if bullets_ok
           else "Use bullets for achievements and keep paragraphs short — recruiters scan, they don't read."),
    )

    # quantification — ≥30% of bullets carry a number/%/currency.
    if bullet_lines:
        quantified = [ln for ln in bullet_lines if _QUANT_RE.search(ln)]
        ratio = len(quantified) / len(bullet_lines)
        quant_ok = ratio >= 0.3
        quant_detail = (
            f"{len(quantified)}/{len(bullet_lines)} bullets are quantified "
            f"({round(ratio * 100)}%). "
            + ("Numbers are the strongest impact signal — good."
               if quant_ok
               else "Add numbers (%, $, counts, latency, users) to at least a third of your bullets.")
        )
    else:
        quant_ok, quant_detail = False, "No bullet points found to quantify."
    add("quantification", "Quantified achievements", quant_ok, 15, quant_detail)

    # dates — one consistent format family.
    month_hits = _MONTH_DATE_RE.findall(text)
    stripped_text = _MONTH_DATE_RE.sub("", text)
    numeric_hits = _NUMERIC_DATE_RE.findall(stripped_text)
    year_hits = _YEAR_RANGE_RE.findall(_NUMERIC_DATE_RE.sub("", stripped_text))
    families = [n for n in (len(month_hits), len(numeric_hits), len(year_hits)) if n >= 2]
    any_dates = bool(month_hits or numeric_hits or year_hits)
    dates_ok = any_dates and len(families) <= 1
    add(
        "dates", "Consistent date format", dates_ok, 10,
        "One consistent date style." if dates_ok
        else ("No dates found — ATS ranking uses employment dates." if not any_dates
              else "Mixed date styles (e.g. 'Jan 2024' and '01/2024') confuse ATS date parsing — pick one."),
    )

    # first-person — resumes are written in implied first person, no I/me/my.
    fp = len(_FIRST_PERSON_RE.findall(text))
    add(
        "first-person", "No first-person pronouns", fp < 5, 5,
        "Implied first person throughout." if fp < 5
        else f"{fp} uses of I/me/my — drop the pronouns ('Led migration…', not 'I led…').",
    )

    # parse-risk — column/table artifacts from extraction.
    pipe_lines = sum(1 for ln in nonempty if ln.count("|") >= 2)
    gap_lines = sum(1 for ln in nonempty if _MULTISPACE_RE.search(ln))
    parse_ok = pipe_lines < 3 and gap_lines < 6
    add(
        "parse-risk", "Single-column, parse-safe layout", parse_ok, 10,
        "No column or table artifacts detected." if parse_ok
        else "Extraction shows column/table artifacts (pipes or wide gaps mid-line) — multi-column layouts scramble ATS reading order; use a single column.",
    )

    # buzzword-stuffing — filler phrases or one keyword repeated >6×.
    lower = text.lower()
    fillers = [p for p in _FILLER_PHRASES if p in lower]
    tokens = [
        t for t in re.findall(r"[a-z][a-z+#.-]{3,}", lower)
        if t not in _STOPWORDS and t not in _STUFFING_EXEMPT
    ]
    common = Counter(tokens).most_common(1)
    stuffed = common[0] if common and common[0][1] > 6 else None
    stuffing_ok = not fillers and stuffed is None
    add(
        "buzzword-stuffing", "No stuffing or filler", stuffing_ok, 10,
        "No filler buzzwords or keyword stuffing." if stuffing_ok
        else "Found: "
        + "; ".join(
            filter(None, [
                ("filler phrases (" + ", ".join(fillers) + ")") if fillers else "",
                (f"'{stuffed[0]}' repeated {stuffed[1]}×") if stuffed else "",
            ])
        )
        + ". Modern ATS treats stuffing as a fraud signal.",
    )

    # weak-language — duty-speak and filler ("responsible for", "leveraged").
    weak_hits = find_weak_resume_phrases(text)
    weak_total = sum(n for _, n in weak_hits)
    weak_ok = weak_total < 3
    add(
        "weak-language", "Verb-first, concrete language", weak_ok, 8,
        "No duty-speak or filler; bullets lead with strong verbs."
        if not weak_hits else
        ("Found " + ", ".join(f"'{p}' ×{n}" for p, n in weak_hits[:4])
         + (". Acceptable in small doses" if weak_ok
            else ". Rewrite verb-first: 'Led X', 'Cut Y by 30%', not what you were 'responsible for'")
         + "."),
    )

    # Structure-aware checks — only when we have the structured resume.
    if structured is not None:
        undated = [w.name or w.position for w in structured.work if not w.startDate]
        add(
            "work-dates", "Every job has dates", not structured.work or not undated, 8,
            "All work entries carry start dates." if not undated
            else f"Missing start dates on: {', '.join(filter(None, undated)) or 'some entries'} — undated experience gets discounted by ATS ranking.",
        )
        has_skills = any(sk.keywords or sk.name for sk in structured.skills)
        add(
            "skills-section", "Dedicated skills section", has_skills, 8,
            "Skills section present — the easiest keyword-match surface."
            if has_skills
            else "No skills section found — ATS keyword matching leans on it heavily.",
        )

    coverage = compute_keyword_coverage(text, jd_text) if jd_text else None
    return AtsReport(score=_score(checks, coverage), checks=checks, keyword_coverage=coverage)


# ── JD keyword extraction (deterministic) ───────────────────────────────────

_REQ_HEADER_RE = re.compile(
    r"^\s*(?:requirements?|qualifications?|must[- ]haves?|what you.ll need|"
    r"what we.re looking for|about you|skills?)\b",
    re.IGNORECASE,
)


def extract_jd_keywords(jd_text: str, limit: int = 18) -> list[str]:
    """Frequency-weighted unigrams + bigrams from the JD, minus stopwords,
    with lines under a Requirements-style header counted double."""
    weights: Counter[str] = Counter()
    in_requirements = False
    for line in jd_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if _REQ_HEADER_RE.match(stripped):
            in_requirements = True
        elif stripped.endswith(":") or (len(stripped) < 60 and stripped.isupper()):
            in_requirements = _REQ_HEADER_RE.match(stripped) is not None
        boost = 2 if in_requirements else 1
        tokens = [
            t for t in (
                # Keep interior punctuation (node.js, ci/cd) and trailing +/#
                # (c++, c#), but strip sentence punctuation off the end.
                raw.rstrip("./-")
                for raw in re.findall(r"[a-z0-9][a-z0-9+#./-]{1,}", stripped.lower())
            )
            if t and t not in _STOPWORDS and not t.isdigit() and len(t) >= 2
        ]
        for t in tokens:
            if len(t) >= 3 or t in ("go", "c#", "r"):
                weights[t] += boost
        for a, b in zip(tokens, tokens[1:]):
            if len(a) >= 3 and len(b) >= 3:
                weights[f"{a} {b}"] += boost
    # Prefer bigrams over their constituent unigrams when both rank.
    ranked = [t for t, _ in weights.most_common(limit * 3)]
    picked: list[str] = []
    for term in ranked:
        if len(picked) >= limit:
            break
        if " " not in term and any(term in bg.split() for bg in picked):
            continue
        picked.append(term)
    return picked


def compute_keyword_coverage(resume_text: str, jd_text: str) -> KeywordCoverage:
    keywords = extract_jd_keywords(jd_text)
    if not keywords:
        return KeywordCoverage(found=[], missing=[], percent=100)
    haystack = re.sub(r"\s+", " ", resume_text.lower())
    found = [k for k in keywords if k in haystack]
    missing = [k for k in keywords if k not in haystack]
    return KeywordCoverage(
        found=found,
        missing=missing,
        percent=round(100 * len(found) / len(keywords)),
    )
