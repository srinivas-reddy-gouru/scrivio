"""Extract plain text from an uploaded resume (PDF, DOCX, or plain text).

The parser is deliberately dumb: no layout reconstruction, no section
detection — the LLM downstream reads the raw text better than any
heuristic re-formatter. The only jobs here are (1) getting text out of
binary containers and (2) failing with an ACTIONABLE message when we
can't (the classic case: a scanned-image PDF with no text layer).
"""
from __future__ import annotations

import io

# ~30k chars ≈ 7-8k tokens: generous for any real resume (even academic
# CVs), small enough to leave prompt room for the JD and instructions.
_MAX_CHARS = 30_000

# Below this many extracted characters a PDF is almost certainly a scan
# (image-only pages) rather than a text document.
_SCAN_SUSPECT_CHARS = 40


class ResumeParseError(ValueError):
    """Raised when a resume can't be turned into usable text. The message
    is shown to the user verbatim — keep it actionable."""


def parse_resume(data: bytes, filename: str) -> str:
    """Return the resume's plain text, capped at _MAX_CHARS."""
    if not data:
        raise ResumeParseError("The uploaded file is empty.")
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        text = _parse_pdf(data)
    elif name.endswith(".docx"):
        text = _parse_docx(data)
    elif name.endswith((".txt", ".md", ".text")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ResumeParseError(
            "Unsupported file type. Upload a PDF, DOCX, or TXT file — or "
            "paste the resume text directly."
        )

    text = text.strip()
    if len(text) < _SCAN_SUSPECT_CHARS:
        raise ResumeParseError(
            "Almost no text could be extracted — this PDF looks like a "
            "scanned image. Paste the resume text directly instead."
        )
    return text[:_MAX_CHARS]


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise ResumeParseError(
                "This PDF is password-protected. Remove the password or "
                "paste the text directly."
            )
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ResumeParseError:
        raise
    except Exception:
        raise ResumeParseError(
            "Could not read this PDF. Try re-exporting it, or paste the "
            "resume text directly."
        )


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception:
        raise ResumeParseError(
            "Could not read this DOCX file. Try re-saving it, or paste the "
            "resume text directly."
        )
